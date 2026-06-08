# -*- coding: utf-8 -*-
"""
HIA 评估报告 Word 导出。
用 python-docx 生成正式 docx:标题页 + 7 章节,中文宋体正文 / 黑体标题。
结构沿用 v0.9 工作流的 7 章节。
"""

from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SONG = "宋体"          # 正文
HEI = "黑体"           # 标题
GREEN = RGBColor(0x1B, 0x6B, 0x3A)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)


def _font(run, name=SONG, size=12, bold=False, color=None):
    """统一设置中文字体(关键:同时设 eastAsia,否则中文回退默认字体)。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _para(doc, text="", size=12, bold=False, color=None, align=None,
          name=SONG, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _font(run, name=name, size=size, bold=bold, color=color)
    return p


def _heading(doc, text, level=1):
    """章节标题:黑体。level=1 为一级章节标题。"""
    size = 15 if level == 1 else 13
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _font(run, name=HEI, size=size, bold=True, color=GREEN)
    return p


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    # 表头
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(h)
        _font(r, name=HEI, size=10.5, bold=True)
    # 数据行
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[j].paragraphs[0].add_run(str(val))
            _font(r, name=SONG, size=10.5)
    return t


def build_report_docx(case, result) -> BytesIO:
    """根据案例与计算结果生成报告,返回内存中的 docx 字节流。"""
    years = case["evaluation_years"]
    mname = result["metric_name"]              # 可避免病例数 / 额外病例数
    cooling = result["direction"] == "cooling"
    cvd = result["cvd_primary"]
    ac = result["ac_primary"]
    sens = result["cvd_sensitivity"]
    accent = GREEN if cooling else RED

    def f1(x):
        return f"{abs(x):.1f}"

    def f2(x):
        return f"{abs(x):.2f}"

    doc = Document()
    # Normal 默认字体设为宋体
    normal = doc.styles["Normal"]
    normal.font.name = SONG
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    # ===== 标题页 =====
    for _ in range(3):
        doc.add_paragraph()
    _para(doc, "上海市城市规划项目", size=26, bold=True, color=accent,
          align=WD_ALIGN_PARAGRAPH.CENTER, name=HEI, space_after=4)
    _para(doc, "健康影响评估专项报告", size=26, bold=True, color=accent,
          align=WD_ALIGN_PARAGRAPH.CENTER, name=HEI, space_after=18)
    _para(doc, case["project_name"], size=15, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    for _ in range(6):
        doc.add_paragraph()
    _para(doc, "同济大学建筑与城市规划学院 · 健康城市实验室", size=13,
          align=WD_ALIGN_PARAGRAPH.CENTER, name=HEI, space_after=4)
    _para(doc, "AI 辅助 · 热暴露—健康影响量化 (HIA) 系统", size=11,
          color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, f"报告日期:{date.today().isoformat()}", size=11,
          color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== 第一章 项目概况 =====
    _heading(doc, "一、项目概况")
    _para(doc, case["background"])
    direction_zh = "降温场景(预期健康效益)" if cooling else "升温场景(预期健康风险)"
    _table(
        doc,
        ["项目要素", "内容"],
        [
            ["项目名称", case["project_name"]],
            ["所在区位", f"{case['city']} · {case['district']} · {case['subdistrict']}"],
            ["干预类型", case["intervention_type_zh"]],
            ["建设规模", f"{case['size_ha']} 公顷"],
            ["受影响人口", f"{case['population']:,} 人({case['population_brief']})"],
            ["评估期", f"{years} 年"],
            ["气候带", "夏热冬冷地区"],
            ["场景方向", direction_zh],
        ],
    )

    # ===== 第二章 暴露变化 =====
    _heading(doc, "二、暴露变化")
    _table(
        doc,
        ["暴露指标", "变化值", "说明"],
        [
            ["地表温度 ΔLST", f"{case['delta_lst_C']:+.1f} °C",
             f"夏季地表温度基线约 {case['summer_lst_baseline_C']} °C"],
            ["近地气温 ΔT_air", f"{case['delta_air_temp_C']:+.2f} °C",
             "由 ΔLST 经经验系数转换,作为健康量化输入"],
        ],
    )
    _para(doc, (
        f"绿地{'新增' if cooling else '损失'}引起地块及周边夏季热环境的"
        f"{'改善' if cooling else '恶化'},近地气温变化 ΔT_air = "
        f"{case['delta_air_temp_C']:+.2f} °C 为后续健康影响量化的核心暴露输入。"
    ))

    # ===== 第三章 健康影响量化结果 =====
    _heading(doc, "三、健康影响量化结果")
    _para(doc, (
        f"采用双暴露-反应关系(ERF)并行估计。在 {years} 年评估期内,"
        f"该项目预计带来心血管疾病 {f1(cvd['total_point'])} 例、"
        f"全因死亡 {f1(ac['total_point'])} 例的{mname}(主分析)。"
    ))
    _para(doc, "主分析与敏感性分析结果(单位:例)", bold=True, space_after=4)
    _table(
        doc,
        ["指标 / ERF", f"每年{mname}", f"累计 {years} 年", "95% CI(累计)"],
        [
            ["心血管(主分析 · Liu 2022)", f2(cvd["per_year_point"]),
             f1(cvd["total_point"]), f"{f1(cvd['total_low'])} – {f1(cvd['total_high'])}"],
            ["全因(主分析 · Liu 2022)", f2(ac["per_year_point"]),
             f1(ac["total_point"]), f"{f1(ac['total_low'])} – {f1(ac['total_high'])}"],
            ["心血管(敏感性 · Lingyan 2017)", f2(sens["per_year_point"]),
             f1(sens["total_point"]), f"{f1(sens['total_low'])} – {f1(sens['total_high'])}"],
        ],
    )
    # 年度累计表(心血管主分析)
    _para(doc, "心血管主分析的逐年累计(示意)", bold=True, space_after=4)
    annual_rows = []
    for y in (1, max(1, years // 2), years):
        annual_rows.append([f"第 {y} 年", f1(cvd["per_year_point"] * y)])
    _table(doc, ["评估年", f"累计{mname}(例)"], annual_rows)

    # ===== 第四章 关键假设与方法说明 =====
    _heading(doc, "四、关键假设与方法说明")
    ep_cvd = result["erf_primary_cvd"]
    ep_ac = result["erf_primary_allcause"]
    es = result["erf_sensitivity_cvd"]
    method_items = [
        f"主分析 ERF:{ep_cvd['label']};心血管 RR={ep_cvd['rr_point']}、"
        f"全因 RR={ep_ac['rr_point']}({ep_cvd['unit']})。",
        f"敏感性 ERF:{es['label']};RR={es['rr_point']}({es['unit']})。",
        f"城市基线死亡率:全因 {result['baseline_allcause_per_1000']}‰、"
        f"心血管 {result['baseline_cvd_per_1000']}‰(上海市卫健委统计公报)。",
        f"热归因比例:取夏季热相关死亡占比 {result['heat_fraction']:.0%}。",
        f"计算式:可避免/额外病例 = 热相关基线病例 ×(1 − RR^ΔT)× 评估年数,"
        f"其中 ΔT_air = {case['delta_air_temp_C']:+.2f} °C。",
    ]
    for it in method_items:
        _para(doc, "· " + it, space_after=4)

    # ===== 第五章 敏感性分析的意义 =====
    _heading(doc, "五、敏感性分析的意义")
    ratio = result["sensitivity_ratio"]
    ratio_txt = f"约为主分析的 {ratio:.1f} 倍" if ratio else "—"
    _para(doc, (
        f"敏感性分析采用更陡峭的 Lingyan 2017 暴露-反应关系作为上界情景,"
        f"其心血管估计{ratio_txt}。该区间用于刻画 ERF 选择带来的不确定性,"
        "真实健康影响大概率落在主分析与敏感性分析估计之间。"
    ))

    # ===== 第六章 局限性 =====
    _heading(doc, "六、局限性")
    limits = [
        "ERF 来自人群流行病学 meta 分析,迁移到本地存在不确定性;",
        "仅量化热暴露的死亡终点,未涵盖发病、住院、心理与社会效益;",
        "ΔLST→ΔT_air 转换与降温幅度基于经验系数,未做精细微气候模拟;",
        "人口为周边估算常住口径,未细分昼夜流动与脆弱人群空间分布;",
        "未考虑空调使用、行为调整等适应行为对暴露的削减;",
        "评估期内人口规模与基线死亡率假定保持恒定;",
        "地块多边形与部分参数为演示用途的合理估计,非实测数据;",
        "结果用于规划决策辅助,不替代正式的环境健康风险评估。",
    ]
    for i, it in enumerate(limits, 1):
        _para(doc, f"{i}. {it}", space_after=4)

    # ===== 第七章 规划建议 =====
    _heading(doc, "七、规划建议")
    if cooling:
        recs = [
            "优先实施并尽量提质:量化结果显示明确的健康效益,建议保障绿地规模与乔木覆盖以稳定降温。",
            "面向脆弱人群布局:在老年人、户外劳动者集中处增设遮荫与饮水等避暑设施,放大健康收益。",
            "建立监测:对地表温度与就诊数据做前后对比,为后续项目积累本地化 ERF 证据。",
        ]
    else:
        recs = [
            "审慎评估与缓解:拆除将带来可量化的健康风险,建议核算是否可避免或缩小拆除范围。",
            "同步补偿降温:就近新增绿地、增加遮荫与喷雾等措施,抵消热岛加剧的短期影响。",
            "错峰与告知:将施工拆除避开盛夏高温,并对周边脆弱人群做高温健康提示。",
        ]
    for i, it in enumerate(recs, 1):
        _para(doc, f"{i}. {it}", space_after=4)

    # ===== 落款 =====
    doc.add_paragraph()
    _para(doc, "报告生成:HIA 智能体(Phase 3.P1+P2+P3+P4)", size=10, color=GREY,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _para(doc, "本结果由本地计算核心生成,用于演示。", size=10, color=GREY,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
