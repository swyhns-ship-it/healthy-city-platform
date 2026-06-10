# -*- coding: utf-8 -*-
"""AI 辅助 HIA 初筛 —— 引擎(纯逻辑,被 views/hia_screen.py 调用)。

对标国内《健康影响评估初筛表》:评估对象文档(PDF/Word)+ 10 个重点问题,
逐条给出「是/不知道/否」研判草案 + 影响路径 + 原文依据 + 证据缺口,供单专家复核。
AI 仅辅助研判,不替代专家判定与签字。

职责:① 文档抽取(pypdf / python-docx);② 调 DeepSeek 出结构化草案(稳健 JSON + 退化重试);
③ 生成填好的初筛表 docx。判断阈值/结论档位等确定性逻辑留在页面与本模块的纯函数里,不交给 LLM。
"""
import json
from io import BytesIO
from datetime import date

import requests

API_URL = "https://api.deepseek.com/chat/completions"
MODEL = "deepseek-chat"
MAX_DOC_CHARS = 40000          # 文档过长时截断(MVP 不做切分/检索)
ANSWERS = ("是", "不知道", "否")

# —— 初筛表的 10 个重点问题(与表格文字一致)+ 每条维度释义(给 LLM 做 grounding)——
QUESTIONS = [
    "可能导致人群传染病和感染性疾病的发生发展。",
    "可能加剧人群重点慢性病的发生发展。",
    "可能增加人群中毒和伤害事件的风险。",
    "可能增加其他突发公共卫生事件的风险。",
    "可能对人口高质量发展带来不利影响。",
    "可能对空气、饮用水、食品和环境卫生等健康环境带来不利影响。",
    "可能对人群健康生活方式、社会心理健康等带来不利影响。",
    "可能对卫生健康投入保障和医疗保险水平带来不利影响。",
    "可能对优质医疗资源合理配置和利用带来不利影响。",
    "可能对医疗卫生服务质量安全和利用、公平性和可及性带来不利影响。",
]
DIM_GUIDE = [
    "传染病/感染性疾病:人口聚集、卫生设施、病媒孳生、给排水、人员流动等引发的传播风险。",
    "重点慢性病:体力活动、饮食环境、空气/噪声等长期暴露对心脑血管、呼吸、代谢等慢病的影响。",
    "中毒与伤害:危化品、交通、生产安全、建筑施工等导致的急性中毒与意外伤害。",
    "其他突发公共卫生事件:群体性事件、食源性、职业健康、灾害次生的公共卫生风险。",
    "人口高质量发展:生育友好、妇幼健康、老龄健康、人口结构与素质相关的不利影响。",
    "健康环境:空气、饮用水、食品安全、环境卫生、土壤等环境健康要素的不利改变。",
    "健康生活方式与心理:出行/锻炼/社交空间、社会心理压力、健康公平的社会决定因素。",
    "卫生投入与医保:公共卫生与医疗投入、医保保障水平的削弱。",
    "优质医疗资源配置:优质医疗资源的可获得性、布局合理性与利用效率。",
    "卫生服务质量与可及性:医疗卫生服务的安全、利用、公平性与可及性。",
]

_SYS = """你是国内「健康影响评估(HIA)初筛」的辅助研判助手,服务对象是评估专家。
任务:阅读「评估对象」文档,针对《健康影响评估初筛表》的 10 个重点问题逐条研判,
判断该规划/政策/工程项目是否「可能」对该维度人群健康带来相关影响。

研判准则(务必遵守):
- 每题给出 answer:仅限「是 / 不知道 / 否」。问题问的是「是否*可能*带来影响」,有合理路径即可判「是」。
- evidence 必须引用文档中的相关原文片段;**若文档没有可支撑的信息,answer 应倾向「不知道」并在 gaps 说明需要专家补充核实什么**。不要凭空编造原文,不要臆断。
- pathway 写出简短影响路径:行动/要素 → 健康决定因素 → 健康结果。
- confidence 为 0–1,反映你对该判断的把握。
- 你是辅助,最终由专家判定;语气中立、克制,不夸大。

只输出一个 JSON 对象(不要任何额外文字、不要 markdown 代码块):
{
 "items": [{"q": 1, "answer": "是|不知道|否", "confidence": 0.0,
            "pathway": "影响路径", "evidence": "文档原文依据(无则空字符串)",
            "gaps": "证据缺口/需专家核实点"}, ... 共 10 条, q 从 1 到 10],
 "summary": "整体研判小结(2–4 句中文)",
 "suggest_level": "很小|轻度|重大"
}
suggest_level 仅为参考建议:多数题为「否」且无重大路径→很小;有少数「是」/中等关注→轻度;
涉及多维度或严重健康风险→重大。"""


def extract_text(name, data):
    """从上传文件抽取纯文本。name 用后缀判类型,data 为 bytes。
    返回 (text, info)。info 含 kind / pages / truncated / error。"""
    info = {"kind": "", "pages": 0, "truncated": False, "error": ""}
    low = (name or "").lower()
    try:
        if low.endswith(".pdf"):
            info["kind"] = "PDF"
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            info["pages"] = len(reader.pages)
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        elif low.endswith(".docx"):
            info["kind"] = "Word"
            from docx import Document
            doc = Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:                       # 含表格文字
                for row in tbl.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            text = "\n".join(parts)
        elif low.endswith(".doc"):
            info["error"] = "旧版 .doc 不支持,请另存为 .docx 或 PDF 后上传。"
            return "", info
        else:
            info["error"] = "仅支持 PDF 或 Word(.docx)。"
            return "", info
    except Exception as e:
        info["error"] = f"文档解析失败:{e}"
        return "", info

    text = (text or "").strip()
    if not text:
        info["error"] = "未能从文档提取到文字(可能是扫描件/图片型 PDF,需 OCR;本工具暂不支持)。"
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS]
        info["truncated"] = True
    return text, info


def _questions_block():
    return "\n".join(f"{i+1}. {q}(维度释义:{DIM_GUIDE[i]})"
                     for i, q in enumerate(QUESTIONS))


def _extract_json(content):
    s = (content or "").strip()
    if s.startswith("```"):
        s = s.strip("`").strip()
        if s[:4].lower() == "json":
            s = s[4:].strip()
    i, j = s.find("{"), s.rfind("}")
    if i >= 0 and j > i:
        s = s[i:j + 1]
    try:
        d = json.loads(s)
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}


def _normalize(data):
    """把模型输出规整成 10 条齐全、字段类型正确的结构;缺失项补「不知道」占位。"""
    by_q = {}
    for it in (data.get("items") or []):
        if not isinstance(it, dict):
            continue
        try:
            q = int(it.get("q"))
        except Exception:
            continue
        ans = str(it.get("answer", "")).strip()
        if ans not in ANSWERS:
            ans = "不知道"
        try:
            conf = float(it.get("confidence", 0))
        except Exception:
            conf = 0.0
        by_q[q] = {
            "q": q, "answer": ans, "confidence": max(0.0, min(1.0, conf)),
            "pathway": str(it.get("pathway", "") or "").strip(),
            "evidence": str(it.get("evidence", "") or "").strip(),
            "gaps": str(it.get("gaps", "") or "").strip(),
        }
    items = []
    for q in range(1, len(QUESTIONS) + 1):
        items.append(by_q.get(q, {"q": q, "answer": "不知道", "confidence": 0.0,
                                  "pathway": "", "evidence": "",
                                  "gaps": "模型未给出该项,请专家自行研判。"}))
    level = str(data.get("suggest_level", "")).strip()
    if level not in ("很小", "轻度", "重大"):
        level = ""
    return {"items": items, "summary": str(data.get("summary", "") or "").strip(),
            "suggest_level": level}


def _call_once(doc_text, project_name, api_key, timeout, temperature, nudge=""):
    user = (f"【评估对象名称】{project_name or '(未填)'}\n"
            f"【初筛表 10 个重点问题】\n{_questions_block()}\n\n"
            f"【评估对象文档全文(可能已截断)】\n{doc_text}\n\n"
            f"请按系统要求,对 10 个问题逐条研判并只输出 JSON。{nudge}")
    r = requests.post(
        API_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"model": MODEL,
              "messages": [{"role": "system", "content": _SYS},
                           {"role": "user", "content": user}],
              "temperature": temperature,
              "response_format": {"type": "json_object"}, "max_tokens": 4000},
        timeout=timeout)
    r.raise_for_status()
    content = r.json()["choices"][0]["message"]["content"] or ""
    return _extract_json(content)


def screen(doc_text, api_key, project_name="", timeout=120):
    """对文档做 10 题初筛研判,返回 {items[10], summary, suggest_level}。
    沿用 llm_agent 的稳健策略:json_object 偶发吐空白 → 升温 + 扰动重试。"""
    attempts = [(0.3, ""), (0.7, "务必输出完整 10 条 items,answer 只用「是/不知道/否」。"),
                (1.0, "用 JSON 回我,items 必须 10 条,evidence 无依据就留空并写 gaps。")]
    data = {}
    for temp, nudge in attempts:
        try:
            data = _call_once(doc_text, project_name, api_key, timeout, temp, nudge)
        except Exception:
            continue
        if data.get("items"):
            break
    return _normalize(data)


# ============ 初筛表 docx 导出 ============
def build_screen_docx(header, items, level, expert_opinion):
    """生成填好的《健康影响评估初筛表》docx(单专家辅助版)。
    header: dict(name/category/dept/submitter/phone/screen_date/method/related_dept)
    items: 经专家复核后的 10 条 [{q, answer, pathway, evidence, gaps, note}]
    level: 健康影响程度(很小/轻度/重大);expert_opinion: 专家意见文本。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    SONG, HEI = "宋体", "黑体"
    GREEN = RGBColor(0x1B, 0x6B, 0x3A)
    GREY = RGBColor(0x55, 0x55, 0x55)

    def font(run, name=SONG, size=10.5, bold=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def para(text="", size=10.5, bold=False, color=None, align=None, name=SONG, after=4):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        font(p.add_run(text), name=name, size=size, bold=bold, color=color)
        return p

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = SONG
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    para("健康影响评估初筛表", size=18, bold=True, name=HEI, color=GREEN,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    # —— 表头信息 ——
    info = [("评估对象名称", header.get("name", "")),
            ("发布/实施类别", header.get("category", "")),
            ("起草/提交部门", header.get("dept", "")),
            ("提交人", header.get("submitter", "")),
            ("电话", header.get("phone", "")),
            ("初筛日期", header.get("screen_date", "")),
            ("初筛方法", header.get("method", "")),
            ("涉及的相关部门", header.get("related_dept", ""))]
    t0 = doc.add_table(rows=0, cols=2)
    t0.style = "Table Grid"
    for k, v in info:
        cells = t0.add_row().cells
        font(cells[0].paragraphs[0].add_run(k), name=HEI, size=10.5, bold=True)
        font(cells[1].paragraphs[0].add_run(str(v or "")), size=10.5)
    doc.add_paragraph()

    # —— 10 题研判表 ——
    para("健康影响评估应重点关注的问题", size=12, bold=True, name=HEI, after=6)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    for j, h in enumerate(["#", "重点关注的问题", "是", "不知道", "否"]):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(c.paragraphs[0].add_run(h), name=HEI, size=10.5, bold=True)
    for it in items:
        cells = t.add_row().cells
        font(cells[0].paragraphs[0].add_run(str(it["q"])), size=10.5)
        font(cells[1].paragraphs[0].add_run(QUESTIONS[it["q"] - 1]), size=10.5)
        for k, label in enumerate(ANSWERS):
            cells[2 + k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            mark = "☑" if it.get("answer") == label else "□"
            font(cells[2 + k].paragraphs[0].add_run(mark), size=11,
                 bold=(it.get("answer") == label))
    doc.add_paragraph()

    # —— 研判依据(AI 辅助 + 专家备注)——
    para("研判依据与影响路径(AI 辅助 · 专家核定)", size=12, bold=True, name=HEI, after=6)
    for it in items:
        para(f"{it['q']}. 判定:{it.get('answer','')}　路径:{it.get('pathway','') or '—'}",
             size=10, bold=True, after=2)
        if it.get("evidence"):
            para(f"   依据原文:{it['evidence']}", size=9.5, color=GREY, after=2)
        if it.get("gaps"):
            para(f"   证据缺口:{it['gaps']}", size=9.5, color=GREY, after=2)
        if it.get("note"):
            para(f"   专家备注:{it['note']}", size=9.5, after=4)

    # —— 结论 ——
    doc.add_paragraph()
    para("评估专家组意见", size=12, bold=True, name=HEI, after=4)
    para(expert_opinion or "(待专家填写)", size=10.5, after=8)
    levels = ["很小", "轻度", "重大"]
    line = "结论:健康影响程度　" + "　".join(
        ("☑" if level == lv else "□") + lv for lv in levels)
    para(line, size=11, bold=True, after=10)

    para("专家组长审定签字:____________　日期:____年__月__日", size=10.5, after=4)
    para("参与专家签字:____________　日期:____年__月__日", size=10.5, after=10)

    para("说明:本表由「AI 辅助 HIA」工具协助生成——AI 仅基于上传文档提供研判草案与原文依据,"
         "判定结论以专家核定与签字为准,AI 不替代专家判定。", size=9, color=GREY)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
